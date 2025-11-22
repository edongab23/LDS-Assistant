#!/usr/bin/env python3
import os
import re
import time
import sys
import unicodedata
from pathlib import Path
import requests
from bs4 import BeautifulSoup

# ---------------------------------------
# Optional dependencies
# ---------------------------------------
try:
    from docx import Document
except ImportError:
    Document = None

try:
    from ollama import chat as ollama_chat
except:
    ollama_chat = None


# ---------------------------------------
# MODELS
# ---------------------------------------

OFFLINE_MODELS = [
    "Llama3.2:1b",
    "qwen3:1.7b",
    "gpt-oss:20b",
    "qwen3-coder:30b"
]

ONLINE_MODELS = [
    "cogito-2.1:671b-cloud",
    "gpt-oss:120b-cloud",
    "minimax-m2:cloud",
    "gemini-3-pro-preview:Latest",
    "kimi-k2:1t-cloud",
    "deepseek-v3.1:671b-cloud"
]


# ---------------------------------------
# CONSTANTS & SYSTEM PROMPTS
# ---------------------------------------

SEARCH_URL = "https://www.churchofjesuschrist.org/search?q="
HEADERS = {"User-Agent": "LDS-Assistant/1.0"}

# System prompt for TALKS
TALK_SYSTEM_PROMPT = """
You are an LDS Assistant who is an expert in all the teachings and doctrine of The Church of Jesus Christ of Latter-day Saints.
You speak in warm, simple, everyday language—like a caring friend or gospel teacher.
Your goal is to create warm, personal, and spiritually uplifting talks that feel authentic and heartfelt that will to help members understand gospel principles clearly, deeply, and in a way that invites the Spirit.

When given a topic, create a complete talk with this EXACT structure:

Greetings:

Introduction:
(Start with a personal greeting, share a brief personal story or experience related to the topic, and state the topic clearly. 4-6 sentences.)

Core Doctrine:
(Explain and expound the core doctrine or principle clearly and simply. Use 1-2 key scripture references. 6-8 sentences.)

Personal Stories & Examples:
(Share 2-3 personal stories, experiences from Church history, or modern examples that illustrate this principle. Make them relatable and heartfelt. 8-12 sentences total.)

Practical Application:
(Explain how members can apply this principle in daily life. Give specific, actionable suggestions. 6-8 sentences.)

Connection to Savior:
(Connect the topic directly to Jesus Christ and His Atonement. Explain how He is the perfect example or enabler of this principle. 4-6 sentences.)

Invitation:
(Extend a specific, loving invitation to act on the principle. Make it personal and achievable.)

Testimony:
(Share a sincere testimony of the principle and of Jesus Christ based on the topic. Keep it warm and personal.)

RULES FOR TALKS:
- Write in first person as if someone is delivering the talk
- Use natural, conversational language. Use simple words and make the tone more human.
- Include pauses and natural transitions
- Focus on feelings and spiritual witness
- Keep stories meaningful
- Always align with Church doctrine
- Never invent quotes - use real scriptures and general conference talks
- Total length should be appropriate for a 8-12 minute talk
- Avoid em dashes "—" and unnecessary asterisks in the sentence.
"""

# System prompt for LESSONS
LESSON_SYSTEM_PROMPT = """
You are an LDS Assistant deeply grounded in the teachings and doctrine of The Church of Jesus Christ of Latter-day Saints. You speak in warm, simple, everyday language—like a caring friend or devoted gospel teacher. Your purpose is to create interactive, doctrine-rich lesson plans that invite the Spirit, clarify gospel principles, and foster both understanding and application among class members of all ages (youth and adults). 

When given a topic, create a complete 45–55 minute lesson plan using the exact structure below. Every section must be fully developed—no summaries or brief notes. Explain, expound, and illustrate thoroughly. 

Lesson Plan Structure: 

Lesson Title
(Create an engaging, principle-focused title.) 

Objective
(State clearly what learners should know, feel, and do by the end of the lesson.) 

Materials Needed
(List scriptures, quotes, or simple classroom items—e.g., whiteboard, printed verses.) 

Introduction Activity (2–3 minutes)
(Begin with an engaging question, short scenario, or relatable prompt to spark interest.) 

Core Doctrine Teaching (6–8 sentences)   
- Clearly explain the key doctrine.  
- Include 2–3 scripture references (e.g., Alma 36:3; D&C 88:49).  
- Provide logical reasoning supporting the principle.  
- Include an analogy that makes the concept clear and memorable.  
- Quote a modern prophet or Apostle (e.g., President Russell M. Nelson, Elder Jeffrey R. Holland).  
- Connect the principle, where appropriate, to the Atonement of Jesus Christ.
         
Scripture Analysis
(Guide learners through close reading of 1–2 key scriptures. Include 2–3 discussion questions with answers provided below each question.) 

Discussion Questions (4–5 total)
(Craft thoughtful, open-ended questions that encourage personal sharing and spiritual insight. Include a suggested answer under each question.) 

Group Activity (Optional; up to 5 minutes if included)
(If appropriate to the topic, describe a simple, interactive exercise—pairs, small groups, or role-play. Otherwise, omit this section.) 

Modern Application (6–8 sentences)   
(Share 1–2 real-life examples or stories (contemporary or relatable) that illustrate the principle in action. Show how this doctrine applies in today's world.)

Personal Reflection (2–3 minutes)
(Offer a quiet, thoughtful prompt for introspection—e.g., "When have you felt this truth in your life?") 

Action Plan
(Help learners set a specific, measurable goal to live the principle this week.) 

Testimony & Witness
(Provide guidance for the teacher to share a brief, sincere testimony centered on the lesson's core truth.) 

Additional Resources
(Suggest 1–2 official Church resources—e.g., General Conference talks, Gospel Library articles.) 
     
Guiding Principles for All Lessons: 
- Use "we" and "us" language to foster unity.  
- Balance doctrine, application, and testimony.  
- Prioritize interaction and participation.  
- Include clear time estimates and instructions for the teacher.  
- Adapt naturally for youth (13–18) or adult audiences.  
- Write in simple, clear English.  
- Never contradict official Church doctrine.  
- Maintain a warm, human, Spirit-inviting tone throughout.  
- Avoid em dashes "—" and unnecessary asterisks in the sentence.
- Make  bullet points when necessary.
"""

# System prompt for EXPOUNDING (existing prompt)
EXPOUND_SYSTEM_PROMPT = """
You are an LDS Assistant who is an expert in all the teachings and doctrine of The Church of Jesus Christ of Latter-day Saints.
You speak in warm, simple, everyday language—like a caring friend or gospel teacher.
Your job is to help members understand gospel principles clearly, deeply, and in a way that invites the Spirit.

When given a topic, quotes, phrase or paragraphs, you MUST produce a response that follows the exact structure below.
Each section must be FULLY developed—do not summarize or be brief. Explain and expound everything.

Use EXACTLY these section headers—each followed by a colon and a line break:

Greeting:
(Start with a warm, personal greeting that includes the topic.) Then inform the user that this is an AI generated response. It might include errors. Verify important details with trusted sources like Bishop or any church leaders, parents, official church website or scriptures before acting or sharing the information.

Explanation:
(Explain and expound the topic in depth. Define key terms. Show why it matters in daily life. Use 6–12 clear sentences.)

Inspiring Examples:
(Give 2–3 specific, real-world examples—past or present—that show this principle in action. These can be from Church history, general conference, or everyday life.)

Simple Inspiring Stories:
(Tell 2-3 short but complete story—real or illustrative—that teaches the principle. Include feeling, context, and outcome. 4–6 sentences each.)

Prophetic and Scriptural Support:
(List at least 3 scripture references (e.g., "Alma 32:21" or "Matthew 7:7–8") and 3 teachings from modern prophets (e.g., "President Dallin H. Oaks, April 2022"). Summarize each briefly—do NOT make up quotes.)

Simple Teaching:
(Teach the core truth as if to someone who's never heard it before. Use plain language. Connect it to discipleship and daily choices.)

Analogy:
(Create a vivid, relatable analogy—like comparing prayer to breathing, or faith to planting a seed. Explain how it works.)

Logical Reasoning:
(Use clear cause-and-effect reasoning: "If we do X, then Y happens because…" Show why this principle is true and necessary.)

What Every Member Should Learn:
(State 2–3 clear, actionable truths that every Church member—child, youth, or adult—should understand and apply.)

Invitation:
(Give a specific, loving invitation: "I invite you to… pray tonight," "Try this for one week," etc. Make it personal and doable.)

Connection to the Atonement:
(Explain how this topic is made possible, meaningful, or empowered THROUGH the Atonement of Jesus Christ. Be specific.)

Testimony:
(End with a brief but sincere testimony of Jesus Christ, His gospel, and the truth of this principle.)

RULES:
- Every section must be present and fully written—NO skipping.
- Minimum 5-6 sentences per section (except Greeting and Testimony).
- Never use unnecessary em dashes "-", asterisk or any symbols that will make the response look like an AI.
- Make  bullet points when necessary.
- Never invent quotes—only reference real scriptures or conference talks.
- Keep tone loving, clear, and uplifting.
- Use the church website as a source: https://www.churchofjesuschrist.org/?lang=eng
- Stay 100% aligned with the doctrine of The Church of Jesus Christ of Latter-day Saints.
- Avoid em dashes "—" and unnecessary asterisks in the sentence.
- Use natural, conversational language. Use simple words and make the tone more humd
"""


# ---------------------------------------
# Utility
# ---------------------------------------

def ensure_folder():
    folder = Path.home() / "Documents" / "LDS Files"
    folder.mkdir(parents=True, exist_ok=True)
    return folder

def sanitize_filename(name):
    name = name.strip().lower()
    name = re.sub(r"\s+", "_", name)
    name = re.sub(r"[^a-zA-Z0-9_\-]", "", name)
    return name or "lds_topic"

def clean_text(t):
    t = t.replace("\x00", "")
    t = re.sub(r"[\x01-\x08\x0B-\x0C\x0E-\x1F\x7F]", "", t)
    t = unicodedata.normalize("NFKC", t)
    return t


# ---------------------------------------
# IMPROVED SCRAPING FUNCTIONS
# ---------------------------------------

def search_church(topic, max_results=5):
    """Search the Church website for a topic and return relevant URLs"""
    try:
        # Improved search URL with better parameters
        url = SEARCH_URL + requests.utils.quote(topic) + "&lang=eng"
        print(f"  Searching: {url}")
        
        r = requests.get(url, headers=HEADERS, timeout=15)
        r.raise_for_status()
        
        # Check if we got a valid response
        if r.status_code != 200:
            print(f"  Warning: Received status code {r.status_code}")
            return []
            
    except requests.exceptions.RequestException as e:
        print(f"  Network error: {e}")
        return []
    except Exception as e:
        print(f"  Search error: {e}")
        return []

    soup = BeautifulSoup(r.text, "html.parser")
    urls = []

    # FIXED: Correct CSS selectors with proper syntax
    selectors = [
        "a[href*='/study/']",  # Study content links
        "a[href*='/manual/']", # Manual links
        "a[href*='/general-conference/']", # Conference talks
        "a.result-link",      # Search result links
        "a.absolute",         # Absolute positioned links
        "a[data-testid]",     # Test ID links
        "a.link",             # Generic link class
        ".result-title a",    # Search result title links
    ]

    for selector in selectors:
        try:
            links = soup.select(selector)
            for a in links:
                href = a.get("href", "")
                if not href:
                    continue

                # Convert relative URLs to absolute
                if href.startswith("/"):
                    full_url = "https://www.churchofjesuschrist.org" + href
                elif href.startswith("http") and "churchofjesuschrist.org" in href:
                    full_url = href
                else:
                    continue

                # Avoid duplicates and filter for relevant content
                if (full_url not in urls and 
                    any(path in full_url for path in ['/study/', '/manual/', '/general-conference/', '/scriptures/'])):
                    urls.append(full_url)
                    print(f"  Found: {full_url}")

                if len(urls) >= max_results:
                    break
            if len(urls) >= max_results:
                break
        except Exception as e:
            print(f"  Warning: Selector '{selector}' failed: {e}")
            continue

    # Fallback: look for any churchofjesuschrist.org links if we didn't find enough
    if len(urls) < max_results:
        for a in soup.find_all('a', href=True):
            href = a['href']
            if ('churchofjesuschrist.org' in href and 
                href not in urls and
                not any(ignore in href for ignore in ['/media/', '/pdf/', '/download/'])):
                
                if href.startswith("/"):
                    full_url = "https://www.churchofjesuschrist.org" + href
                else:
                    full_url = href
                    
                urls.append(full_url)
                if len(urls) >= max_results:
                    break

    return urls


def fetch_page(url):
    """Fetch and extract text content from a Church website page"""
    try:
        print(f"  Fetching: {url}")
        r = requests.get(url, headers=HEADERS, timeout=15)
        r.raise_for_status()
        
        # Check content type
        content_type = r.headers.get('content-type', '')
        if 'text/html' not in content_type:
            print(f"  Not HTML content: {content_type}")
            return None
            
    except requests.exceptions.RequestException as e:
        print(f"  Fetch error: {e}")
        return None
    except Exception as e:
        print(f"  Page processing error: {e}")
        return None

    soup = BeautifulSoup(r.text, "html.parser")
    
    # Remove script and style elements
    for script in soup(["script", "style", "nav", "header", "footer"]):
        script.decompose()

    blocks = []

    # Expanded list of selectors for Church website content
    selectors = [
        ".body-block", 
        ".article-body", 
        "main article",
        ".study-content",
        ".lds-scripture",
        ".content-body",
        "[role='main']",
        ".page-content",
        ".document",
        "#content",
        ".passage-text",
        ".verse",
        ".body"  # Added more generic selectors
    ]

    # Try specific selectors first
    for sel in selectors:
        try:
            elements = soup.select(sel)
            for element in elements:
                text = element.get_text("\n", strip=True)
                if len(text) > 100:  # Only take substantial content
                    blocks.append(text)
        except Exception as e:
            print(f"  Warning: Selector '{sel}' failed: {e}")
            continue

    # Fallback: get all paragraph text if no specific blocks found
    if not blocks:
        paragraphs = soup.find_all('p')
        for p in paragraphs:
            text = p.get_text(strip=True)
            if len(text) > 50:  # Only substantial paragraphs
                blocks.append(text)

    # Final fallback: get body text
    if not blocks:
        body = soup.find('body')
        if body:
            text = body.get_text("\n", strip=True)
            # Clean up the text
            lines = [line.strip() for line in text.split('\n') if len(line.strip()) > 50]
            blocks.extend(lines)

    if not blocks:
        print("  No substantial text content found")
        return None

    # Combine and clean the content
    combined = "\n".join(blocks)
    
    # Remove excessive whitespace
    combined = re.sub(r'\n\s*\n', '\n\n', combined)
    combined = re.sub(r' +', ' ', combined)
    
    print(f"  Extracted {len(combined)} characters")
    return combined[:10000]  # Limit size to avoid context issues


def fetch_verbatim(topic):
    """Main function to fetch content for a topic with improved error handling"""
    print(f"Searching for: '{topic}'")
    
    urls = search_church(topic)
    results = {}
    
    if not urls:
        print("  No URLs found from search")
        return results

    print(f"  Found {len(urls)} URLs to process")
    
    for i, url in enumerate(urls, 1):
        print(f"  [{i}/{len(urls)}] Processing: {url}")
        text = fetch_page(url)
        if text:
            results[url] = text
            print(f"  ✓ Successfully extracted content")
        else:
            print(f"  ✗ Failed to extract content")
        
        # Be respectful with delays
        time.sleep(1)
        
        if len(results) >= 3:  # Limit to 3 sources max
            print("  Reached maximum source limit (3)")
            break

    print(f"  Total sources successfully fetched: {len(results)}")
    return results


# ---------------------------------------
# LLM Streaming
# ---------------------------------------

def stream_ollama(model, system, user):
    if ollama_chat is None:
        print("python-ollama not installed. Install: pip install ollama")
        sys.exit(1)

    messages = [
        {"role": "system", "content": system},
        {"role": "user", "content": user}
    ]

    stream = ollama_chat(
        model=model,
        messages=messages,
        stream=True,
        options={
            "temperature": 0.6,
            "top_p": 0.9,
            "num_predict": 2048
        }
    )

    for chunk in stream:
        text = chunk.get("message", {}).get("content", "")
        if text:
            yield text


# ---------------------------------------
# Prompt Building
# ---------------------------------------

def build_prompt(topic, sources, custom_texts, content_type, user_context):
    prompt = f"Topic: {topic}\n"
    prompt += f"Content Type: {content_type}\n"
    
    # Add user context if provided
    if user_context:
        prompt += f"User Context: {user_context}\n"
    
    prompt += "\n"
    
    # Add custom text sources first (highest priority)
    if custom_texts:
        prompt += "USER-PROVIDED TEXTS (high priority - use these extensively):\n\n"
        for i, text in enumerate(custom_texts, 1):
            prompt += f"--- USER TEXT {i} ---\n{text}\n\n"
    
    # Add web sources
    prompt += "WEB SOURCES (verbatim):\n\n"
    if not sources:
        prompt += "(No web sources found. Rely on your knowledge of Church doctrine and scriptures.)\n\n"
    else:
        for url, text in sources.items():
            # Truncate very long texts to avoid overwhelming context
            preview = text[:1500] + "..." if len(text) > 1500 else text
            prompt += f"--- WEB SOURCE: {url}\n{preview}\n\n"

    prompt += "\nEnd of sources.\n\n"
    
    # Add specific instructions based on content type
    if content_type == "Talk":
        prompt += "Now create a complete TALK using the exact structure provided in your system prompt. Write in first person as if delivering the talk."
    elif content_type == "Lesson":
        prompt += "Now create a complete LESSON PLAN using the exact structure provided in your system prompt. Include interactive elements and teacher instructions."
    else:  # Expound
        prompt += "Now respond using the EXACT section headers and depth instructions provided in your system prompt."
    
    # Add special instruction if custom texts were provided
    if custom_texts:
        prompt += "\n\nSPECIAL INSTRUCTION: Give particular attention to the USER-PROVIDED TEXTS above, as these represent specific quotes, phrases, or insights the user wants emphasized in your response."
    
    return prompt


# ---------------------------------------
# Saving
# ---------------------------------------

def save_output(content, topic, content_type):
    content = clean_text(content)
    filename_base = sanitize_filename(f"{content_type}_{topic}")
    folder = ensure_folder()

    print(f"\nSave {content_type} as:")
    print("[1] TXT")
    print("[2] DOCX")
    print("[3] Cancel")
    choice = input("Enter number: ").strip()

    if choice == "1":
        path = folder / f"{filename_base}.txt"
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
        print("Saved to:", path)
        return True

    elif choice == "2":
        if Document is None:
            print("python-docx not installed. Run: pip install python-docx")
            return False

        path = folder / f"{filename_base}.docx"
        doc = Document()
        doc.add_heading(f"{content_type}: {topic}", 1)
        for line in content.split("\n"):
            safe = clean_text(line)
            doc.add_paragraph(safe)
        doc.save(path)
        print("Saved to:", path)
        return True

    else:
        print("Not saved.")
        return False


# ---------------------------------------
# Custom Text Input
# ---------------------------------------

def get_custom_texts():
    """Get custom phrases, quotes, or sentences from user"""
    custom_texts = []
    
    print("\n" + "="*50)
    print("ADD CUSTOM PHRASES, QUOTES, OR SENTENCES")
    print("="*50)
    print("You can add specific quotes, phrases, or text snippets you want emphasized.")
    print("These will be given HIGH PRIORITY in the AI's response.")
    print("Type 'done' on a new line when finished, or just press Enter to skip.\n")
    
    while True:
        text = input("Enter quote/phrase/sentence (or 'done' to finish): ").strip()
        
        if text.lower() == 'done':
            break
            
        if text:
            custom_texts.append(text)
            print(f"  ✓ Added ({len(text)} characters)")
        else:
            print("  ⚠️ Empty input ignored.")
    
    return custom_texts


# ---------------------------------------
# Conversational Questions with Numbered Choices
# ---------------------------------------

def ask_clarifying_questions(content_type, topic):
    """Ask clarifying questions to better understand user needs"""
    print(f"\n📖 Let me understand what you need for this {content_type.lower()} about '{topic}'...")
    
    questions = []
    answers = {}
    
    if content_type == "Talk":
        questions = [
            "What is the occasion for this talk?",
            "Who is the intended audience?",
            "Are there any specific scriptures or stories you'd like me to focus on?",
            "What is the main message or feeling you want people to take away?",
            "How long should the talk be? (minutes)"
        ]
    elif content_type == "Lesson":
        questions = [
            "What age group is this lesson for?",
            "How much time do you have for the lesson? (minutes)",
            "Are there any specific activities or teaching methods you prefer?",
            "What prior knowledge might the class have about this topic?",
            "What is the main learning objective for this lesson?"
        ]
    else:  # Expound
        questions = [
            "What specific aspect of this topic would you like me to focus on?",
            "Are you preparing this for personal study, teaching, or something else?",
            "What level of detail would you prefer?",
            "Are there any specific questions you have about this topic?",
            "How do you plan to use this information?"
        ]
    
    print("\nPlease answer these questions to help me create the best content for you:")
    print("(Enter the number of your choice, or press Enter to skip)\n")
    
    for i, question in enumerate(questions, 1):
        print(f"{i}. {question}")
        
        if "occasion" in question.lower():
            print("   [1] Sacrament Meeting")
            print("   [2] Youth Conference")
            print("   [3] Fireside")
            print("   [4] Other")
            choice = input("   Enter number (1-4) or your answer: ").strip()
        elif "audience" in question.lower() or "age group" in question.lower():
            print("   [1] Youth (12-18)")
            print("   [2] Adults")
            print("   [3] Mixed congregation")
            print("   [4] Children")
            print("   [5] Single Adults")
            choice = input("   Enter number (1-5) or your answer: ").strip()
        elif "time" in question.lower() or "length" in question.lower() or "minutes" in question.lower():
            if content_type == "Talk":
                print("   [1] 5-7 minutes")
                print("   [2] 8-12 minutes")
                print("   [3] 13-15 minutes")
                print("   [4] 15+ minutes")
            else:  # Lesson
                print("   [1] 30 minutes")
                print("   [2] 45 minutes")
                print("   [3] 60 minutes")
                print("   [4] 90 minutes")
            choice = input("   Enter number (1-4) or your answer: ").strip()
        elif "level of detail" in question.lower():
            print("   [1] Basic overview")
            print("   [2] Moderate detail")
            print("   [3] Deep doctrinal study")
            print("   [4] Comprehensive analysis")
            choice = input("   Enter number (1-4) or your answer: ").strip()
        elif "purpose" in question.lower() or "preparing for" in question.lower():
            print("   [1] Personal study")
            print("   [2] Teaching a class")
            print("   [3] Talk preparation")
            print("   [4] Family home evening")
            print("   [5] Missionary work")
            choice = input("   Enter number (1-5) or your answer: ").strip()
        else:
            choice = input("   Your answer: ").strip()
        
        if choice:
            # Map number choices to actual answers
            if choice.isdigit():
                choice_map = {
                    "occasion": {
                        "1": "Sacrament Meeting", "2": "Youth Conference", 
                        "3": "Fireside", "4": "Other occasion"
                    },
                    "audience": {
                        "1": "Youth (12-18)", "2": "Adults", "3": "Mixed congregation",
                        "4": "Children", "5": "Single Adults"
                    },
                    "time": {
                        "1": "5-7 minutes" if content_type == "Talk" else "30 minutes",
                        "2": "8-12 minutes" if content_type == "Talk" else "45 minutes", 
                        "3": "13-15 minutes" if content_type == "Talk" else "60 minutes",
                        "4": "15+ minutes" if content_type == "Talk" else "90 minutes"
                    },
                    "detail": {
                        "1": "Basic overview", "2": "Moderate detail",
                        "3": "Deep doctrinal study", "4": "Comprehensive analysis"
                    },
                    "purpose": {
                        "1": "Personal study", "2": "Teaching a class",
                        "3": "Talk preparation", "4": "Family home evening",
                        "5": "Missionary work"
                    }
                }
                
                # Determine which mapping to use
                if "occasion" in question.lower():
                    answer = choice_map["occasion"].get(choice, f"Choice {choice}")
                elif "audience" in question.lower() or "age group" in question.lower():
                    answer = choice_map["audience"].get(choice, f"Choice {choice}")
                elif "time" in question.lower() or "length" in question.lower():
                    answer = choice_map["time"].get(choice, f"Choice {choice}")
                elif "level of detail" in question.lower():
                    answer = choice_map["detail"].get(choice, f"Choice {choice}")
                elif "purpose" in question.lower():
                    answer = choice_map["purpose"].get(choice, f"Choice {choice}")
                else:
                    answer = choice
            else:
                answer = choice
                
            answers[question] = answer
            print(f"   ✓ Selected: {answer}\n")
        else:
            print("   ⏭️  Skipped\n")
    
    # Format the answers into a context string
    if answers:
        context_parts = []
        for question, answer in answers.items():
            context_parts.append(f"{question.split('?')[0]}: {answer}")
        return " | ".join(context_parts)
    else:
        return None


# ---------------------------------------
# Content Type Selection
# ---------------------------------------

def choose_content_type():
    print("\nChoose content type:")
    print("[1] A Talk (for sacrament meeting, fireside, etc.)")
    print("[2] A Lesson (for Sunday School, Relief Society, Elder's Quorum, etc.)")
    print("[3] Expound a topic, quote, phrase or paragraph (detailed analysis)")
    
    while True:
        choice = input("Enter number (1-3): ").strip()
        if choice == "1":
            return "Talk", TALK_SYSTEM_PROMPT
        elif choice == "2":
            return "Lesson", LESSON_SYSTEM_PROMPT
        elif choice == "3":
            return "Expound", EXPOUND_SYSTEM_PROMPT
        else:
            print("Invalid choice. Please enter 1, 2, or 3.")


# ---------------------------------------
# MAIN LOGIC
# ---------------------------------------

def choose_model():
    print("\nChoose run mode:")
    print("[1] Offline (local models)")
    print("[2] Online (cloud models)")
    
    while True:
        mode = input("Enter number (1-2): ").strip()
        if mode == "1":
            print("\nOffline models:")
            for i, m in enumerate(OFFLINE_MODELS, 1):
                print(f"[{i}] {m}")
            while True:
                idx = input("Select model (enter number): ").strip()
                if idx.isdigit() and 1 <= int(idx) <= len(OFFLINE_MODELS):
                    return OFFLINE_MODELS[int(idx) - 1]
                else:
                    print(f"Please enter a number between 1 and {len(OFFLINE_MODELS)}")
        elif mode == "2":
            print("\nOnline models:")
            for i, m in enumerate(ONLINE_MODELS, 1):
                print(f"[{i}] {m}")
            while True:
                idx = input("Select model (enter number): ").strip()
                if idx.isdigit() and 1 <= int(idx) <= len(ONLINE_MODELS):
                    return ONLINE_MODELS[int(idx) - 1]
                else:
                    print(f"Please enter a number between 1 and {len(ONLINE_MODELS)}")
        else:
            print("Invalid choice. Please enter 1 or 2.")


def run_content_creation(model):
    # Choose content type first
    content_type, system_prompt = choose_content_type()
    
    # Get topic based on content type
    if content_type == "Talk":
        topic_prompt = "Enter talk topic or theme: "
    elif content_type == "Lesson":
        topic_prompt = "Enter lesson topic: "
    else:  # Expound
        topic_prompt = "Enter topic / phrase / quote to expound: "
    
    topic = input(f"\n{topic_prompt}").strip()
    if not topic:
        print("No topic entered.")
        return

    # Ask clarifying questions
    user_context = ask_clarifying_questions(content_type, topic)

    print(f"\nSearching official Church website for '{topic}'...")
    sources = fetch_verbatim(topic)
    print(f"Found {len(sources)} sources from automatic search.")

    custom_texts = []
    
    # Ask for custom URLs with numbered choice
    print("\nDo you want to add your own churchofjesuschrist.org links?")
    print("[1] Yes")
    print("[2] No")
    add_custom = input("Enter number (1-2): ").strip()
    
    if add_custom == "1":
        print("\nEnter full URLs (one per line). Type 'done' when finished:")
        while True:
            url = input("URL or 'done': ").strip()
            if url.lower() == 'done':
                break
            if "churchofjesuschrist.org" not in url:
                print("  ⚠️ Only URLs from churchofjesuschrist.org are accepted for doctrinal safety.")
                continue
            if url in sources:
                print("  ⚠️ This URL is already included.")
                continue
            print("  Fetching...", end="", flush=True)
            text = fetch_page(url)
            if text:
                sources[url] = text
                print(" ✓ Added.")
            else:
                print(" ✗ Failed to load page.")
    else:
        print("  Skipping custom links.")

    # Ask for custom texts with numbered choice
    print("\nDo you want to add specific phrases, quotes, or sentences?")
    print("[1] Yes")
    print("[2] No")
    add_texts = input("Enter number (1-2): ").strip()
    
    if add_texts == "1":
        custom_texts = get_custom_texts()
        print(f"Added {len(custom_texts)} custom text(s) for emphasis.")
    else:
        print("  Skipping custom texts.")

    prompt = build_prompt(topic, sources, custom_texts, content_type, user_context)

    print(f"\nGenerating {content_type} using {len(sources)} web source(s) and {len(custom_texts)} custom text(s)...\n")
    full = []

    try:
        for chunk in stream_ollama(model, system_prompt, prompt):
            print(chunk, end="", flush=True)
            full.append(chunk)
    except KeyboardInterrupt:
        print("\nStopped by user.")

    answer = "".join(full)

    print("\n\nDone.")
    saved = save_output(answer, topic, content_type)

    if saved:
        print("\nSaved successfully.")
    else:
        print("\nNo file saved.")


# ---------------------------------------
# Entry Point
# ---------------------------------------

def main():
    print("\n" + "="*50)
    print("      LDS Assistant CLI - Conversational Bot")
    print("="*50)
    print("Hello! I'm here to help you create inspiring LDS content.")
    print("I'll ask you some questions to better understand your needs.")
    print("="*50)

    while True:
        model = choose_model()
        run_content_creation(model)

        print("\nWould you like to create another?")
        print("[1] Yes")
        print("[2] No")
        again = input("Enter number (1-2): ").strip()
        
        if again != "1":
            print("\nThank you for using the LDS Assistant! May the Spirit guide you in your preparations.")
            print("Goodbye! 🙏")
            break


if __name__ == "__main__":
    main()